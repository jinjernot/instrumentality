from docx.shared import Pt, Inches
from docx import Document

from app.routes.ds_tool.core.table import table_column_widths
from app.routes.ds_tool.core.footer import add_footer

def excel_to_word(df, new_df, header_value):

    # Create a Word document
    doc = Document("/home/garciagi/frame/template.docx")
    
    # Add header with the extracted value
    section = doc.sections[0]  # Get the first section (assuming only one section)
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_run = header_paragraph.add_run(header_value)
    header_run.bold = True  # Make the text bold
    header_run.font.size = Pt(16)  # Set the font size to 16pt

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Series Values")
    run.font.size = Pt(12)
    run.bold = True

    # Add the DataFrame 'df' as a table to the Word document
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
    table_column_widths(table, (Inches(2), Inches(5.5),))
    
    # Populate the table with 'df' data
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = column  # Set column headers
        for j in range(df.shape[0]):
            value = df.iloc[j, i]
            cell = table.cell(j + 1, i)
            cell.text = str(value)
            # Set font size for each cell in the table
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)  # Set font size to 6 points
            # Reduce spacing within cell
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Set font size and spacing for table header
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)  # Set font size to 6 points
                run.font.bold = True
        # Add spacing to table header
        cell.paragraphs[0].paragraph_format.space_before = Pt(6)
        cell.paragraphs[0].paragraph_format.space_after = Pt(6)

    # Set table width to match entire page width
    table.autofit = False
    section = doc.sections[-1]
    table_width = section.page_width - section.left_margin - section.right_margin
    table.width = table_width

    doc.add_paragraph()

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("SKU Values")
    run.font.size = Pt(12)
    run.bold = True

    for i in range(1, new_df.shape[1]):
        # Add a paragraph break between tables
        #doc.add_paragraph()

        # Add a new table for the current column
        table = doc.add_table(rows=new_df.shape[0] + 1, cols=2,  style='Table Grid')
        table_column_widths(table, (Inches(2), Inches(5.5),))

        # Populate the table with the values from the first and current column
        for j in range(new_df.shape[0]):
            # Set column headers
            table.cell(0, 0).text = new_df.columns[0]
            table.cell(0, 1).text = new_df.columns[i]

            # Set values for current row
            table.cell(j + 1, 0).text = str(new_df.iloc[j, 0])  # First column value
            table.cell(j + 1, 1).text = str(new_df.iloc[j, i])  # Current column value

            # Set font size for each cell in the table
            for cell in table.rows[j + 1].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)  # Set font size to 6 points
                        
                # Reduce spacing within cell
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        # Set font size and spacing for table header
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Set font size to 6 points
                    run.font.bold = True
            # Add spacing to table header
            cell.paragraphs[0].paragraph_format.space_before = Pt(6)
            cell.paragraphs[0].paragraph_format.space_after = Pt(6)

        # Set table width to match entire page width
        table.autofit = False
        table.width = table_width

        # Insert a line break after each table
        doc.add_paragraph()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)  # Set left margin to 0.5 inches
        section.right_margin = Inches(0.5)  # Set right margin to 0.5 inches
        section.top_margin = Inches(0.5)  # Set top margin to 0.5 inches
        section.bottom_margin = Inches(0.5)  # Set bottom margin to 0.5 inches

    add_footer(doc)

    # Save Word file with the same name as the Excel file
    word_filename = '/home/garciagi/frame/ds.docx'
    doc.save(word_filename)
