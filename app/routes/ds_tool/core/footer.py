from docx.shared import Pt

def add_footer(doc):
    # Add a new section with a footer
    section = doc.sections[-1]
    new_footer = section.footer
    new_footer.paragraphs[0].clear()  # Clear existing content if any

    # Add the disclaimer text to the footer
    footer_paragraph = new_footer.paragraphs[0]
    #footer_paragraph.add_run("Disclaimer:").bold = True
    disclaimer_run = footer_paragraph.add_run("\nThis document contains confidential information and is intended only for the product development content review. Do not disseminate, distribute, or copy this document or any of its contents. You cannot use or forward the file without the Product Management team consent or Hewlett-Packard Development Company, L.P. permission. The information contained herein is subject to change without notice. HP shall not be liable for technical or editorial errors or omissions contained herein.")
    disclaimer_run.font.size = Pt(6)