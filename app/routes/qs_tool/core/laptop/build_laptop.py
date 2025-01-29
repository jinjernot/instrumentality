from app.routes.qs_tool.core.laptop.tech_specs.tech_specs import tech_specs_section
from app.routes.qs_tool.core.laptop.overview.overview import overview_section
from app.routes.qs_tool.core.laptop.tables.tables import table_section
from app.routes.qs_tool.core.format.format import format_document

from docx import Document
from zipfile import ZipFile
from config import QS_DOCX_FILE_PATH, QS_ZIP_FILE_PATH


imgs_path = "/home/garciagi/qs/imgs/"
#imgs_path = "./imgs/"

def createdocx(file):
    """Table Secion"""
    
    # Variables
    doc = Document()

    # Quickspecs sections
    overview_section(doc, file)
    tech_specs_section(doc, file)
    table_section(doc, file)

    format_document(doc, file, imgs_path)
    docx_file = QS_DOCX_FILE_PATH
    #docx_file = 'quickspecs.docx'

    doc.save(docx_file)

    # Convert DOCX to PDF using docx2pdf
    #convert(docx_file)

    # Create a zip file and add specific files to it
    zip_file_name = QS_ZIP_FILE_PATH
    #zip_file_name = 'quickspecs.zip'

    with ZipFile(zip_file_name, 'w') as zipf:
        zipf.write(docx_file, arcname='quickspecs.docx')