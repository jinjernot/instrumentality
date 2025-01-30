from app.routes.qs_tool.core.format.table import table_column_widths
from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.format.hr import *

from docx.enum.text import WD_BREAK
from docx.shared import Inches, RGBColor

def docking_section(doc, df):
    """Docking Table"""

    try:
        # Define possible title variations
        docking_title = ["Docking (sold separately)", "Docking"]

        # Check if the DataFrame contains a docking section
        if not any(df.iloc[:, 1].isin(docking_title)):
            return  # Exit early if no docking data exists

        # Insert title
        insert_title(doc, "Docking (Sold Separately)")

        for index, row in df.iterrows():
            if row[1] in docking_title:
                # Add a table with 2 columns
                table = doc.add_table(rows=1, cols=2)

                # Set column widths
                table_column_widths(table, [Inches(3), Inches(5)])

                # Populate the table
                for i in range(index + 1, len(df)):
                    if df.iloc[i, 0] == "Container Name":
                        break  # Stop at the next section

                    # Add row to table
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(df.iloc[i, 0])
                    row_cells[1].text = str(df.iloc[i, 1])

                    # Bold column 0 text
                    if row_cells[0].paragraphs:
                        run = row_cells[0].paragraphs[0].add_run()
                        run.bold = True
                        run.text = row_cells[0].text
                        row_cells[0].paragraphs[0].clear()  # Clear original text
                        row_cells[0].paragraphs[0].add_run(run.text).bold = True

                # Remove the first (empty) row from the table
                if len(table.rows) > 1:
                    table.rows[0]._element.getparent().remove(table.rows[0]._element)

                # Add paragraph break after table
                doc.add_paragraph()

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)
        
        # Insert page break
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        error_msg = f"An error occurred: {e}"
        print(error_msg)

        error_paragraph = doc.add_paragraph()
        error_run = error_paragraph.add_run(error_msg)
        error_run.bold = True
        error_run.font.color.rgb = RGBColor(255, 0, 0)

        return str(e)
