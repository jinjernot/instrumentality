from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.format.hr import *
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor
import pandas as pd
import re

def system_unit_section(doc, file):
    """System Unit table"""
    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only System Unit', engine='openpyxl')

        # Add title: SYSTEM UNIT
        insert_title(doc, "SYSTEM UNIT")

        start_col_idx = 0
        end_col_idx = 1
        start_row_idx = 4
        end_row_idx = 42

        # Find the footnotes row index (if exists)
        footnotes_index = df[df.eq('Footnotes').any(axis=1)].index.tolist()
        if footnotes_index:
            end_row_idx = footnotes_index[0] - 1  # Stop before the footnotes row

        # Extract only the necessary rows
        data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
        data_range = data_range.dropna(how='all')

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=num_rows, cols=num_cols)

        # Define column widths
        column_widths = (Inches(3), Inches(5))

        # Set column widths
        table_column_widths(table, column_widths)

        for row_idx in range(num_rows):
            for col_idx in range(num_cols):
                value = data_range.iat[row_idx, col_idx]
                cell = table.cell(row_idx, col_idx)
                if not pd.isna(value):
                    cell.text = str(value)

        # Bold the first column
        for row in table.rows:
            if row.cells[0].paragraphs and row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].font.bold = True

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                               
        doc.add_paragraph()

        # Process footnotes if they exist
        if footnotes_index:
            footnotes_index = footnotes_index[0]  # Assuming there's only one "Footnotes" row
            footnotes_data = df.iloc[footnotes_index + 1:].dropna(how='all')  # Drop rows with all NaN values
            
            footnotes = [" - ".join(map(str, row.dropna().tolist())) for _, row in footnotes_data.iterrows() if row.dropna().tolist()]
            
            paragraph = doc.add_paragraph()
            pattern = re.compile(r"\[(\d+)\]")  # Match [x] where x is a number

            for index, data in enumerate(footnotes):
                if "Container Name" in data or "Wireless WAN" in data:
                    continue
                
                cleaned_data = pattern.sub(r"\1.", data)  # Replace [x] with x.
                run = paragraph.add_run(cleaned_data)
                run.font.color.rgb = RGBColor(0, 0, 153)  # Set font color to blue
                
                if index < len(footnotes) - 1 and footnotes[index + 1].strip():
                    run.add_break(WD_BREAK.LINE)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)

def table_column_widths(table, widths):
    """Set the column widths for a table."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
