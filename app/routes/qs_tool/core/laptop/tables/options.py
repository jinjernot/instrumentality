from app.routes.qs_tool.core.format.table import table_column_widths
from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.format.hr import *
from docx.shared import Inches

from docx.enum.text import WD_BREAK
import pandas as pd

def options_section(doc, file):
    """Options QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Options', engine='openpyxl')

        # Add title: Options
        insert_title(doc, "OPTIONS")

        start_col_idx = 0
        end_col_idx = 2
        start_row_idx = 1
        end_row_idx = 299

        data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
        data_range = data_range.dropna(how='all')  # Drop rows with all NaN values

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=0, cols=num_cols)  # Start with an empty table

        # Populating table cells with data (dynamically stop at empty row)
        for row_idx in range(num_rows):
            row = data_range.iloc[row_idx]
            if row.isna().all():  # Stop if the entire row is empty
                break

            # Add a row to the table
            table.add_row()

            # Populate the row cells with data
            for col_idx in range(num_cols):
                value = row[col_idx]
                cell = table.cell(row_idx, col_idx)
                if not pd.isna(value):
                    cell.text = str(value)

                    # Bold the first row
                    if row_idx == 0:
                        cell.paragraphs[0].runs[0].font.bold = True
                        
        table_column_widths(table, (Inches(2), Inches(4), Inches(2)))

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)
