import pandas as pd

from docx.enum.text import WD_BREAK

from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.blocks.table import *
from app.routes.qs_tool.core.format.hr import *

def network_section(doc, file, html_file):
    """Network QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Network', engine='openpyxl')

        # Add title: Networking
        insert_title(doc, "NETWORKING / COMMUNICATION", html_file)

        # Add table
        insert_table(doc, df, html_file)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)