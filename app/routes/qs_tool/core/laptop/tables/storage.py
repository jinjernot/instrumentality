import pandas as pd

from docx.enum.text import WD_BREAK

from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.blocks.table import *
from app.routes.qs_tool.core.format.hr import *

def storage_section(doc, file, html_file):
    """Storage QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Storage', engine='openpyxl')

        # Add title: Storage 
        insert_title(doc, "STORAGE", html_file)

        # Add table
        insert_table(doc, df, html_file)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)