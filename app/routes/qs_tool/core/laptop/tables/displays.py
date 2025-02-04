from app.routes.qs_tool.core.blocks.paragraph import *
from app.routes.qs_tool.core.blocks.title import *
from app.routes.qs_tool.core.blocks.table import *
from app.routes.qs_tool.core.format.hr import *
import pandas as pd

from docx.enum.text import WD_BREAK

def displays_section(doc, file):
    """Displays QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Displays', engine='openpyxl')

        # Add title: Displays
        insert_title(doc, "DISPLAYS")
              
        paragraph = doc.add_paragraph()

        run = paragraph.add_run("Actual brightness will be lower with touchscreen or HP Sure View.")
        run.font.color.rgb = RGBColor(0, 0, 153)
        paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run("Availability may vary by country")
        run.font.color.rgb = RGBColor(0, 0, 153)
        paragraph.add_run().add_break(WD_BREAK.LINE)

        # Add table
        insert_table(doc, df)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)