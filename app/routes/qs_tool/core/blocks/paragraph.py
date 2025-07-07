from docx.enum.text import WD_BREAK
import re
from docx.shared import Pt
from docx.shared import RGBColor
from app.routes.qs_tool.core.format.hr import *

def insert_paragraph(doc, df, iloc_row, iloc_column):
    """
    Insert a paragraph into the Word document.
    """
    data = df.iloc[iloc_row, iloc_column]
    paragraph = doc.add_paragraph()
    paragraph.add_run(data)

def process_footnotes(doc, footnotes):
    """
    Add footnotes to the Word document with blue font color,
    while replacing [x] with x. (number followed by a period).
    """
    if not footnotes:
        return

    paragraph = doc.add_paragraph()
    pattern = re.compile(r"\[(\d+)\]")

    for index, data in enumerate(footnotes):
        cleaned_data = pattern.sub(r"\1.", data)

        run = paragraph.add_run(cleaned_data)
        run.font.color.rgb = RGBColor(0, 0, 153)
        
        if index < len(footnotes) - 1:
            run.add_break(WD_BREAK.LINE)

def insert_error(doc, error_message):
    """
    Insert an error message into the Word document with red font color.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Error: {error_message}")
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True
    run.add_break(WD_BREAK.LINE)

def insert_list(doc, df, start_value):
    """
    Insert a list into the Word document.
    """
    if start_value not in df.iloc[:, 1].tolist():
        insert_error(doc, f"'{start_value}' not found in DataFrame.")
        return
    
    start_index = df.index[df.iloc[:, 1] == start_value].tolist()[0]
    
    # Check if both "Certification and Compliance" and "Service and Support" exist
    has_certification = "Certification and Compliance" in df.iloc[:, 1].tolist()
    has_service = "Service and Support" in df.iloc[:, 1].tolist()

    # If the start_value is one of the special cases, determine the end of the list
    if start_value in ["Service and Support", "Certification and Compliance"]:
        # If "Certification and Compliance" exists, it's always the last section
        if has_certification and start_value == "Certification and Compliance":
             items = df.iloc[start_index:, 1].tolist()
        elif has_certification and start_value == "Service and Support":
            # If we start at "Service and Support" but "Certification" also exists,
            # end the list before "Certification and Compliance"
            next_value_indices = df.iloc[start_index:, 1][df.iloc[start_index:, 1] == "Certification and Compliance"].index.tolist()
            if not next_value_indices:
                # This case should ideally not happen if has_certification is true, but as a fallback:
                items = df.iloc[start_index:, 1].tolist()
            else:
                next_value_index = next_value_indices[0]
                items = df.iloc[start_index:next_value_index, 1].tolist()
        else: #This will cover the case where only "Service and Support" is the last section
            items = df.iloc[start_index:, 1].tolist()


    else:
        # Original logic for other sections
        next_value_indices = df.iloc[start_index:, 1][df.iloc[start_index:, 1] == 'Value'].index.tolist()
        if not next_value_indices:
            insert_error(doc, f"'Value' not found after '{start_value}'")
            return
        next_value_index = next_value_indices[0]
        items = df.iloc[start_index:next_value_index, 1].tolist()
        
    # Identify footnotes based on the content starting with '['
    footnotes_index = next(
        (i for i, item in enumerate(items) if item.startswith("[")),
        len(items)
    )
    non_footnotes = items[:footnotes_index]
    footnotes = items[footnotes_index:]
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(start_value.upper())
    
    paragraph = doc.add_paragraph()
    run.font.size = Pt(12)
    run.bold = True
    run.add_break(WD_BREAK.LINE)

    for index, data in enumerate(non_footnotes[1:], start=1):
        pattern = re.compile(r"\[(\d+)\]")

        split_data = pattern.split(data)
        matches = pattern.findall(data)

        for i, text_part in enumerate(split_data):
            if i % 2 == 0:
                run = paragraph.add_run(text_part)
            else:
                sup_run = paragraph.add_run(text_part)
                sup_run.font.superscript = True
                sup_run.font.size = Pt(9)
        
        if index < len(non_footnotes) - 1:
            run.add_break(WD_BREAK.LINE)
    
    run.add_break(WD_BREAK.LINE)
    process_footnotes(doc, footnotes)

    insert_horizontal_line(doc.add_paragraph(), thickness=3)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def insert_footnote(doc, df, iloc_range, iloc_column):
    """
    Insert a footnote into the Word document.
    """
    footnote = df.iloc[iloc_range, iloc_column].tolist()
    paragraph = doc.add_paragraph()
    for index, note in enumerate(footnote):
        run = paragraph.add_run(note)
        run.font.color.rgb = RGBColor(0, 0, 153)
        if index < len(footnote) - 1:
            run.add_break(WD_BREAK.LINE)