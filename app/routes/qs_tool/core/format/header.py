from app.routes.qs_tool.core.format.hr import *

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

import pandas as pd

def header(doc,file):
    """Generate the Word Header"""

    # Get product name
    #df = pd.read_excel(file, sheet_name = 'Callouts') 
    df = pd.read_excel(file.stream, sheet_name='Callouts', engine='openpyxl')

    prod_name = df.columns[1]
    header = doc.sections[0].header
    
    # Create header
    paragraph = header.paragraphs[0]
    p = paragraph._p 
    p.getparent().remove(p)
    header_table = header.add_table(rows=1, cols=2, width=Inches(8))
    header_table.columns[0].width = Inches(4)
    header_table.columns[1].width = Inches(4)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    
    # Set the header title
    header_paragraph = header_table.cell(0, 0).paragraphs[0]
    header_run = header_paragraph.add_run("QuickSpecs")
    header_run.font.size = Pt(27)

    # Add the product name to the header
    header_paragraph = header_table.cell(0, 1).paragraphs[0]
    header_run2 = header_paragraph.add_run(prod_name)
    header_run2.font.size = Pt(14)
    header_run2.font.bold = True
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    insert_horizontal_line(header.add_paragraph(), thickness=30)

    # Add a break line at the end of the header
    header.add_paragraph()