import json

from app.routes.qs_tool.core.format.header import header
from app.routes.qs_tool.core.format.footer import footer

from docx.shared import Pt

def read_bold_words_from_json(json_file):
    """Read bold words from a JSON file"""
    with open(json_file, 'r') as f:
        data = json.load(f)
        return data.get('bold_words', [])

def set_margins(doc):
    """Set document margins"""
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)

def set_default_font(doc):
    """Set default font for the document"""
    styles = doc.styles
    default_style = styles['Normal']
    font = default_style.font
    font.name = 'HP Simplified'
    font.size = Pt(10)

def apply_bold_font(doc, bold_words):
    """
    Apply bold font to specific words in the document"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for word in bold_words:
                if word in run.text:
                    index = run.text.find(word)
                    if index != -1 and (index == 0 or run.text[index - 1] == '\n'):
                        word_end_index = index + len(word)
                        if word_end_index == len(run.text) or run.text[word_end_index] == '\n':
                            run.bold = True

def format_document(doc, file, imgs_path):
    """Apply formatting to the document"""
    bold_words = read_bold_words_from_json('/home/garciagi/qs/app/core/format/bold_words.json')
    #bold_words = read_bold_words_from_json('app/core/format/bold_words.json')
    header(doc, file)
    footer(doc, imgs_path)
    set_margins(doc)
    set_default_font(doc)
    apply_bold_font(doc, bold_words)

    # Apply cell spacing to all tables
    for table in doc.tables:
        table.style.paragraph_format.space_before = Pt(0)
        table.style.paragraph_format.space_after = Pt(0)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.paragraph_format.space_before = Pt(0)
                    paragraph.style.paragraph_format.space_after = Pt(0)
